# loader.py 核心流程：文档解析器（md/pdf/docx/html → 纯文本 + 图片 URL）
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

try:
    from loguru import logger
except Exception:
    import logging
    logger = logging.getLogger(__name__)


@dataclass
class Document:
    """解析后的统一文档格式。"""
    text: str                                # 全文纯文本
    source: str                              # 来源文件名/URL
    doc_title: str                           # 文档标题
    images: List[str] = field(default_factory=list)  # 图片 URL 列表
    sections: List[dict] = field(default_factory=list)  # [{title, text}] 章节信息
    pages: Dict[int, str] = field(default_factory=dict)  # PDF 页码 → 文本


class DocumentLoader:
    """解析 md/pdf/docx/html 成统一 Document 格式。"""

    def load(self, file_path: str) -> Document:
        ext = Path(file_path).suffix.lower()
        if ext == ".md":
            return self._load_md(file_path)
        elif ext == ".pdf":
            return self._load_pdf(file_path)
        elif ext == ".docx":
            return self._load_docx(file_path)
        elif ext in (".html", ".htm"):
            return self._load_html(file_path)
        else:
            raise ValueError(f"不支持的格式: {ext}（支持 md/pdf/docx/html）")

    # ===== Markdown =====
    def _load_md(self, path: str) -> Document:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        source = Path(path).name
        doc_title = self._extract_md_title(text, source)
        images = self._extract_md_images(text)
        sections = self._extract_md_sections(text)
        return Document(text=text, source=source, doc_title=doc_title,
                        images=images, sections=sections)

    def _extract_md_title(self, text: str, fallback: str) -> str:
        match = re.match(r'^#\s+(.+)$', text, re.MULTILINE)
        return match.group(1).strip() if match else fallback

    def _extract_md_images(self, text: str) -> List[str]:
        # ![alt](url) 格式
        urls = re.findall(r'!\[.*?\]\((https?://[^\s)]+)\)', text)
        # HTML <img src="url">
        urls += re.findall(r'<img[^>]+src=["\'](https?://[^"\']+)["\']', text)
        return list(set(urls))

    def _extract_md_sections(self, text: str) -> List[dict]:
        sections = []
        current_title = ""
        current_lines = []
        in_code_fence = False
        for line in text.split("\n"):
            # 跟踪 ``` 代码块状态——块内的 # 开头行不是标题
            if re.match(r'^```', line):
                in_code_fence = not in_code_fence
                current_lines.append(line)
                continue
            if not in_code_fence and re.match(r'^#{1,3}\s+', line):
                if current_title or current_lines:
                    sections.append({"title": current_title, "text": "\n".join(current_lines).strip()})
                current_title = re.sub(r'^#{1,3}\s+', '', line).strip()
                current_lines = []
            else:
                current_lines.append(line)
        if current_title or current_lines:
            sections.append({"title": current_title, "text": "\n".join(current_lines).strip()})
        return sections

    # ===== PDF =====
    def _load_pdf(self, path: str) -> Document:
        from pypdf import PdfReader
        reader = PdfReader(path)
        source = Path(path).name
        pages = {}
        all_text = []
        images = []

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            pages[i + 1] = page_text  # 页码从 1 开始
            all_text.append(page_text)
            # 提取 link annotation URL
            if "/Annots" in page:
                for annot_ref in page["/Annots"]:
                    annot = annot_ref.get_object()
                    if annot.get("/Subtype") == "/Link" and "/A" in annot:
                        action = annot["/A"]
                        if action.get("/S") == "/URI" and "/URI" in action:
                            uri = str(action["/URI"])
                            if uri.startswith("http"):
                                images.append(uri)

        doc_title = Path(path).stem
        return Document(
            text="\n\n".join(all_text),
            source=source,
            doc_title=doc_title,
            images=list(set(images)),
            pages=pages,
        )

    # ===== DOCX =====
    def _load_docx(self, path: str) -> Document:
        import docx
        doc = docx.Document(path)
        source = Path(path).name
        paragraphs = []
        images = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                # 提取 hyperlink
                for rel in para.part.rels.values():
                    if "hyperlink" in str(rel.reltype):
                        if rel.target_url and rel.target_url.startswith("http"):
                            images.append(rel.target_url)

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        doc_title = Path(path).stem

        # DOCX 按标题样式分节
        sections = []
        current_title = ""
        current_lines = []
        for para in doc.paragraphs:
            if para.style and "Heading" in str(para.style.name):
                if current_title or current_lines:
                    sections.append({"title": current_title, "text": "\n".join(current_lines)})
                current_title = para.text.strip()
                current_lines = []
            elif para.text.strip():
                current_lines.append(para.text.strip())
        if current_title or current_lines:
            sections.append({"title": current_title, "text": "\n".join(current_lines)})

        return Document(
            text=full_text,
            source=source,
            doc_title=doc_title,
            images=list(set(images)),
            sections=sections,
        )

    # ===== HTML =====
    def _load_html(self, path: str) -> Document:
        with open(path, "r", encoding="utf-8") as f:
            raw = f.read()
        # 去标签
        text = re.sub(r'<script[^>]*>.*?</script>', '', raw, flags=re.DOTALL)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()

        # 提取标题
        title_match = re.search(r'<title>(.*?)</title>', raw, re.IGNORECASE | re.DOTALL)
        doc_title = title_match.group(1).strip() if title_match else Path(path).name

        # 提取图片 URL
        images = re.findall(r'<img[^>]+src=["\'](https?://[^"\']+)["\']', raw)
        images += re.findall(r'<img[^>]+src=["\']([^"\']+)["\']', raw)

        source = Path(path).name
        return Document(text=text, source=source, doc_title=doc_title, images=list(set(images)))
